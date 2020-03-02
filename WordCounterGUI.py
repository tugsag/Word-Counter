import tkinter as tk
from tkinter import filedialog
import string
import re
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import BytesIO
from docx import Document

m = tk.Tk()
m.title('Text Analysis')
m.geometry("750x450")
fileChosen=False
filename=''
string=''

## Functions
##---------------
def black():
    print('black')

def test():
    print('this is insignificant')

def chooseFile():
    global string
    global fileChosen
    global filename
    filename = filedialog.askopenfilename()
    if filename=='':
        filedisp["text"] = "Please choose a .txt, .pdf, or .docx file before continuing."
    else:
        filedisp["text"] = "Chosen File: " + filename
        if ".txt" in filename:
            fileChosen = True
            with open(filename, 'r') as f:
                f1 = f.read()
                string = f1
        elif ".docx" in filename:
            fileChosen = True
            doc = Document(filename)
            text = ''
            for d in doc.paragraphs:
                text+=d.text

            string=text
        else:
            fileChosen = True
            manager = PDFResourceManager()
            retstr = BytesIO()
            layout = LAParams(all_texts=True)
            device = TextConverter(manager, retstr, laparams=layout)
            filepath = open(filename, 'rb')
            interpreter = PDFPageInterpreter(manager, device)

            for page in PDFPage.get_pages(filepath, check_extractable=True):
                interpreter.process_page(page)

            text = retstr.getvalue()

            filepath.close()
            device.close()
            retstr.close()

            try:
                text = text.decode()
            except AttributeError:
                pass

            string = text
        ##writeFile = tk.Label(m, text="Chosen File: " + filename)
        ##writeFile.pack()

def charcountDisp():
    if fileChosen==False:
        filedisp["text"] = "Please choose a valid file before continuing."
    else:
        cc = len(string)
        ccdisp["text"] = "Character Count: " + str(cc)
        ##charcountDisp = tk.Label(m, text="Character Count: " + str(cc))
        ##charcountDisp.pack()

def wordcountDisp():
    if fileChosen==False:
        filedisp["text"] = "Please choose a valid file before continuing."
    else:
        stripped = re.sub(r'[^\w\s]','',string)
        processed = stripped.lower()
        indWords = processed.split()
        wordCount = len(indWords)
        wcdisp["text"] = "Word Count: " + str(wordCount)
        ##wordcountDisp = tk.Label(m, text="Word Count: " + str(wordCount))
        ##wordcountDisp.pack()

def uniquewordsDisp():
    if fileChosen==False:
        filedisp["text"] = "Please choose a valid file before continuing."
    else:
        stripped = re.sub(r'[^\w\s]','',string)
        processed = stripped.lower()
        indWords = processed.split()
        wordList = []
        for word in indWords:
            if word not in wordList:
                wordList.append(word)

        uwdisp["text"] = "Unique Words: " + str(len(wordList))
        ##uwDisp = tk.Label(m, text="Unique Words: " + str(len(wordList)))
        ##uwDisp.pack()

def charcountwospace():
    if fileChosen==False:
        filedisp["text"] = "Please choose a valid file before continuing."
    else:
        stripped = re.sub("\s+", "", string)
        ccwsdisp["text"] = "Character Count Without Space: " + str(len(stripped))
        ##ccws = tk.Label(m, text="Character Count Without Space: " + str(len(stripped)))
        ##ccws.pack()

def topTen():
    if fileChosen==False:
        filedisp["text"] = "Please choose a valid file before continuing."
    else:
        stripped = re.sub(r'[^\w\s]','',string)
        processed = stripped.lower()
        indWords = processed.split()
        wordList = []
        countList = []
        topTenListWord = []
        topTenListCount = []
        for word in indWords:
            if word not in wordList:
                wordList.append(word)

        for word in wordList:
            numberOfOcc = indWords.count(word)
            countList.append(numberOfOcc)

        for i in range(10):
            maximum = max(countList)
            indexOfMax = countList.index(maximum)
            word = wordList[indexOfMax]
            topTenListWord.append(word)
            topTenListCount.append(maximum)
            countList.remove(maximum)
            wordList.remove(word)

        ttdisp["text"] = "Top Ten: " + str(topTenListWord) + "\n" + "Top Ten Occurrence: " + str(topTenListCount)
        ##tt = tk.Message(m, text="Top Ten: " + str(topTenListWord) + "\n" +
        ##                "Top Ten Occurrence: " + str(topTenListCount))
        ##tt.pack()

def convertPDFtoTXT():
    newName = filename[:-4]
    if fileChosen==False:
        filedisp["text"] = "Please choose a .pdf file before continuing."
    elif ".pdf" not in filename:
        filedisp['text'] = "Please choose a .pdf file before continuing."
        return
    else:
        f=open(newName + ".txt", 'w')
        f.write(string)
        filedisp["text"] = "Success! Find the .txt file in the same directory."
        f.close()

def convertDOCXtoTXT():
    newName = filename[:-5]
    if fileChosen==False:
        filedisp["text"] = "PLease choose a .docx file before continuing."
    elif ".docx" not in filename:
        filedisp["text"] = "PLease choose a .docx file before continuing."
    else:
        f=open(newName + ".txt", 'w')
        f.write(string)
        filedisp["text"] = "Success. Find the .txt file in the same directory."
        f.close()

def convertTXTtoDOCX():
    newName=filename[:-4]
    if fileChosen == False:
        filedisp["text"] = "Please choose a .txt file before continuing."
    elif ".txt" not in filename:
        filedisp["text"] = "Please choose a .txt file before continuing."
    else:
        document=Document()
        document.add_paragraph(string)
        document.save(newName + '.docx')
        filedisp['text'] = "Success. Find the .docx file in the same directory."

## Buttons
##----------------

menubar = tk.Menu(m)
m.config(menu=menubar)
filemenu = tk.Menu(menubar, tearoff=0)
menubar.add_cascade(label="File", menu=filemenu)
filemenu.add_command(label="Open", command=chooseFile)
filemenu.add_command(label="Exit", command=m.destroy)


charcountButton = tk.Button(m, text='Character Count', width=25, command=charcountDisp)
charcountButton.grid(row=0, column=0)

wordcountButton = tk.Button(m, text="Word Count", width=25, command=wordcountDisp)
wordcountButton.grid(row=1, column=0)

uniquewordsButton = tk.Button(m, text="Unique Words", width=25, command=uniquewordsDisp)
uniquewordsButton.grid(row=2, column=0)

ccwsButton = tk.Button(m, text="Char Count W/O Space", width=25, command=charcountwospace)
ccwsButton.grid(row=3, column=0)

ttButton = tk.Button(m, text="Top Ten", height=10, width=25, command=topTen)
ttButton.grid(row=4, column=0)

convertPDFtoTXTButton = tk.Button(m, text="Convert from.pdf to .txt", width=25, command=convertPDFtoTXT)
convertPDFtoTXTButton.grid(row=5, column=0)

convertDOCXtoTXTButton = tk.Button(m, text="Convert from .docx to.txt", width=25, command=convertDOCXtoTXT)
convertDOCXtoTXTButton.grid(row=6, column=0)

convertTXTtoDOCXButton = tk.Button(m, text="Convert from .txt to .docx", width=25, command=convertTXTtoDOCX)
convertTXTtoDOCXButton.grid(row=7, column=0)

## Displays
##------------------

filedisp = tk.Label(m)
filedisp.grid(row=8, column=1)

ccdisp = tk.Label(m)
ccdisp.grid(row=0, column=1)

wcdisp = tk.Label(m)
wcdisp.grid(row=1, column=1)

uwdisp = tk.Label(m)
uwdisp.grid(row=2, column=1)

ccwsdisp = tk.Label(m)
ccwsdisp.grid(row=3, column=1)

ttdisp = tk.Label(m)
ttdisp.grid(row=4, column=1)

m.mainloop()
